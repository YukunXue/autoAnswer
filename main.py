import subprocess
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import re
from docx import Document
import difflib
# encoding:utf-8
import requests
import sys
import json
import base64

from urllib.request import urlopen
from urllib.request import Request
from urllib.error import URLError
from urllib.parse import urlencode
from urllib.parse import quote_plus

import requests
import base64

# client_id 为官网获取的AK， client_secret 为官网获取的SK
API_KEY = ''
SECRET_KEY= ''
TOKEN_URL = 'https://aip.baidubce.com/oauth/2.0/token'
OCR_URL = "https://aip.baidubce.com/rest/2.0/ocr/v1/accurate_basic"

# 函数：使用ADB截取手机屏幕截图
def take_screenshot(file_number):
    subprocess.run(['adb', 'shell', 'screencap', '-p', f'/sdcard/screenshot{file_number}.png'])
    subprocess.run(['adb', 'pull', f'/sdcard/screenshot{file_number}.png', f'screenshot{file_number}.png'])


def fuzzy_match(text1, text2):
    # 使用 SequenceMatcher 计算字符串相似度

    matcher = difflib.SequenceMatcher(None, text1, text2)
    similarity = matcher.ratio()
    return similarity >= 0.5

# 裁剪函数
def crop_image(image_path, left, top, right, bottom):
    image = Image.open(image_path)
    cropped_image = image.crop((left, top, right, bottom))
    return cropped_image

# 预处理函数
def preprocess_image(image_path):
    # 打开图像
    image = Image.open(image_path)
    
    # 去除噪声
    image = image.crop((0, 810, 1216 , 1800))
    image = image.filter(ImageFilter.MedianFilter(size=3))
    
    # 调整对比度和亮度
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(1.5)
    
    # 二值化处理
    image = image.convert('L')
    threshold = 200
    image = image.point(lambda p: p > threshold and 255)

    # 返回预处理后的图像
    return image

# 函数：使用OCR识别图片中的文字
def ocr_image(image_path):
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image, lang='chi_sim')
    return text

# 匹配问题函数
def match_question(text):
    # 匹配问题部分，匹配到问题后的选项部分
    match = re.search(r'问题:(.+?)([A-E].+)', text, re.DOTALL)
    if match:
        question = match.group(1).strip()[:6]  # 只保留前6个字符
        options = match.group(2).strip()    # 提取选项部分
        return question, options
    return None, None

def take_screenshot_v2():
    subprocess.run(['adb', 'shell', 'screencap', '-p', '/sdcard/screenshot.png'])
    subprocess.run(['adb', 'pull', '/sdcard/screenshot.png', 'screenshot.png'])

def match_question_doc(text):
    match = re.search(r'(\d+)\.问题：(.+?)答案：([A-E])', text, re.DOTALL)
    if match:
        question_number = match.group(1).strip()  # 提取问题编号
        question_content = match.group(2).strip()  # 提取问题内容
        correct_answer = match.group(3).strip()  # 提取正确答案
        return question_number, question_content, correct_answer
    return None, None, None

def find_matching_answer(question_text, document_path):
    document = Document(document_path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]  # 提取非空段落文本并去除首尾空白
    for i in range(len(paragraphs)):
        if fuzzy_match(question_text, paragraphs[i]):
            # 找到匹配的问题后，寻找该段最底部的答案
            for j in range(i+1, len(paragraphs)):
                if paragraphs[j].startswith('答案：'):
                    answer_match = re.search(r'答案：([A-E])', paragraphs[j])
                    if answer_match:
                        return answer_match.group(1)
    return None

"""
    获取token
"""
def fetch_token():
    params = {'grant_type': 'client_credentials',
              'client_id': API_KEY,
              'client_secret': SECRET_KEY}
    post_data = urlencode(params)
    post_data = post_data.encode('utf-8')
    req = Request(TOKEN_URL, post_data)
    try:
        f = urlopen(req, timeout=5)
        result_str = f.read()
    except URLError as err:
        print(err)
    result_str = result_str.decode()


    result = json.loads(result_str)

    if ('access_token' in result.keys() and 'scope' in result.keys()):
        if not 'brain_all_scope' in result['scope'].split(' '):
            print ('please ensure has check the  ability')
            exit()
        return result['access_token']
    else:
        print ('please overwrite the correct API_KEY and SECRET_KEY')
        exit()


"""
    读取文件
"""
def read_file(image_path):
    f = None
    try:
        f = open(image_path, 'rb')
        return f.read()
    except:
        print('read image file fail')
        return None
    finally:
        if f:
            f.close()

"""
    调用远程服务
"""
def request(url, data):
    req = Request(url, data.encode('utf-8'))
    has_error = False
    try:
        f = urlopen(req)
        result_str = f.read()
        result_str = result_str.decode()
        return result_str
    except  URLError as err:
        print(err)


def test_api():
    url_req = f"{TOKEN_URL}?grant_type=client_credentials&client_id={API_KEY}&client_secret={SECRET_KEY}"
    
    payload = ""
    headers = {
        'Content-Type': 'application/json',
        'Accept': 'application/json'
    }
    
    response = requests.request("POST", url_req, headers=headers, data=payload)
    
    print(response.text)
    # 将 JSON 字符串解析为 Python 字典
    data = json.loads(response.text)
    
    # 提取你需要的字段
    access_token = data.get('access_token')
    expires_in = data.get('expires_in')
    session_key = data.get('session_key')
    refresh_token = data.get('refresh_token')
    scope = data.get('scope')

    # 打印提取的字段
    print("Access Token:", access_token)
    print("Expires In:", expires_in)
    print("Session Key:", session_key)
    print("Refresh Token:", refresh_token)
    print("Scope:", scope)

def parse_response(api_response):
    question_description = ""
    options_answers ={}
    
    current_option = None
    for item in api_response['words_result']:
        words = item['words']
        if '问题：' in words:
            question_description = words.split('问题：')[-1].strip()
        if re.match(r'^[A-E]$', words):  # 判断是否是选项
            current_option = words
        else:
            # 如果当前选项存在，则将当前文字添加到该选项的答案中
            if current_option:
                options_answers[current_option] = options_answers.get(current_option, "") + words.strip()
                current_option = None
            else :
                match = re.match(r'^([A-E])(.+)$', words)
                if match:
                    current_option, answer = match.groups()
                    options_answers[current_option] = answer.strip()
                    current_option = None

    return options_answers, question_description

def parse_doc(questions):
    document_path = '1.docx'
    document = Document(document_path)
    current_options = {}
    return_val = []
    
    k = 0 #find it
    k1 = 0 
    i = 0
    for paragraph in document.paragraphs:
        if i >= 3:
            i = 0
            
        if '问题：' in paragraph.text:
            question_description = paragraph.text.split('问题：')[-1].strip()
            match = re.match(r'^(.{8})', question_description)
            
            if k == 1 :
                k1 = 1
                
            if match :
                text_quest = match.group()
                if text_quest == questions:
                    k = 1
                    print("yes")
                    print(k)

                    
            i = i + 1
            continue
        
        if i == 1 :
            match = re.search(r'([A-E])\.(.+)', paragraph.text)
            #print(match)
            if match:
                option, text = match.groups()
                #print(option)
                #print(text)
                current_options[option] = text.strip()
            else:
                #print("无法识别的行:", paragraph.text.strip())
                i = i+1

        
        if i == 2 and '答案：' in paragraph.text:
            i = i +1
            match = re.search(r'答案：([A-EYN\|]+)', paragraph.text)

            if match:
                answer = match.group(1)
                #print(answer)
                if "|" in answer:  # 如果答案中包含 | 表示是多选
                    answer = answer.split("|")  # 将多选的答案拆分为列表
                else:
                    answer = [answer]  # 否则将单选的答案放入列表中
                #print(answer)  # 输出: ['A', 'B', 'C']
            if current_options:
                for key, value in current_options.items():
                    if key in answer:
                        if k == 1:
                            return_val.append(value)
                            print("no")
                            
            current_options = {}
        
        if k1 == 1:
            return return_val
    
    return return_val 


#201989
# 主程序
def main():
    preprocessed_image = preprocess_image(f'screenshot14.png')
    preprocessed_image.save(f'screenshot_pre14.png')    
    
    # 获取access token
    token = fetch_token()

    # 拼接通用文字识别高精度url
    image_url = OCR_URL + "?access_token=" + token

    #text = ""

    ## 读取测试图片
    #file_content = read_file('./screenshot_pre5.png')

    ## 调用文字识别服务
    #result = request(image_url, urlencode({'image': base64.b64encode(file_content)}))

    ## 解析返回结果
    #result_json = json.loads(result)
    #print(result_json)
    #for words_result in result_json["words_result"]:
    #    text = text + words_result["words"]

    # 二进制方式打开图片文件
    f = open('./screenshot_pre11.png', 'rb')
    img = base64.b64encode(f.read())

    params = {"image":img}
    headers = {'content-type': 'application/x-www-form-urlencoded'}
    response = requests.post(image_url, data=params, headers=headers)
    if response:
        print(response.json())
    options_answers,  question_description = parse_response(response.json())
    match = re.match(r'^(.{8})', question_description)
    text = match.group()
 
    real_val = []
    real_val = parse_doc(text)
    
    sel_val = []
    print(question_description)  
    print(options_answers)
    print(real_val)
    #print(text)
    if real_val:
        for key, value in options_answers.items():
            if value in real_val:
                sel_val.append(key)
                print(key)
    
if __name__ == "__main__":
    main()
